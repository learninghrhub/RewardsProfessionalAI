import io
from datetime import datetime
from typing import Dict, Tuple

import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st

try:
    from docx import Document
except Exception:
    Document = None

st.set_page_config(
    page_title="RewardAI Professional",
    page_icon="💠",
    layout="wide",
    initial_sidebar_state="expanded",
)

# -------------------------------
# Styling
# -------------------------------
st.markdown(
    """
<style>
    .main .block-container {padding-top: 1.2rem; padding-bottom: 2rem;}
    .ra-hero {background: linear-gradient(135deg,#153c3c 0%,#0f5959 55%,#16a7a7 100%); padding: 26px 30px; border-radius: 22px; color: white; margin-bottom: 18px;}
    .ra-hero h1 {font-size: 2.1rem; margin-bottom: 0.3rem;}
    .ra-hero p {font-size: 1.02rem; opacity: 0.95; margin: 0;}
    .insight-box {border-left: 5px solid #16a7a7; background: #f3fbfb; padding: 14px 18px; border-radius: 12px; margin: 12px 0;}
    .privacy-box {background:#f8fafc; border:1px solid #dbe4e6; padding:13px 15px; border-radius:12px; font-size:0.9rem;}
    .small-note {font-size:0.85rem; color:#667085;}
    .risk-high {color:#b42318; font-weight:700;}
    .risk-medium {color:#b54708; font-weight:700;}
    .risk-low {color:#027a48; font-weight:700;}
    div[data-testid="stMetricValue"] {font-size: 1.28rem;}
</style>
""",
    unsafe_allow_html=True,
)

# -------------------------------
# Sample Data
# -------------------------------

def sample_employee_data() -> pd.DataFrame:
    rows = [
        ["E001", "Ahmed Ali", "Finance", "Finance Manager", "G10", 30000, "Exceeds", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E002", "Sara Khan", "HR", "HR Manager", "G10", 41500, "Meets", "Eligible", "No", "Medium", "Saudi Arabia", "SAR"],
        ["E003", "Omar Fahad", "Operations", "Operations Manager", "G11", 57000, "Partially Meets", "Eligible", "No", "Low", "Saudi Arabia", "SAR"],
        ["E004", "Maya Thomas", "IT", "Cybersecurity Specialist", "G9", 25000, "Outstanding", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E005", "Khalid Saeed", "Sales", "Sales Manager", "G11", 43000, "Exceeds", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E006", "Noura Faisal", "HR", "Total Rewards Specialist", "G8", 20500, "Exceeds", "Eligible", "No", "Medium", "Saudi Arabia", "SAR"],
        ["E007", "Bilal Ahmed", "Finance", "Financial Analyst", "G7", 17800, "Meets", "Eligible", "No", "Low", "Saudi Arabia", "SAR"],
        ["E008", "Reem Hassan", "IT", "Data Analyst", "G8", 17500, "Outstanding", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E009", "John Mathew", "Operations", "Site Supervisor", "G7", 28500, "Meets", "Eligible", "No", "Low", "Saudi Arabia", "SAR"],
        ["E010", "Fatima Noor", "Sales", "Account Executive", "G6", 12600, "Meets", "Eligible", "No", "Medium", "Saudi Arabia", "SAR"],
        ["E011", "Rami Youssef", "IT", "Cloud Engineer", "G9", 29000, "Exceeds", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E012", "Huda Mansour", "Finance", "Finance Director", "G12", 76000, "Outstanding", "Eligible", "Yes", "Medium", "Saudi Arabia", "SAR"],
        ["E013", "Ali Rehman", "HR", "HR Officer", "G6", 11800, "Meets", "Eligible", "No", "Medium", "Saudi Arabia", "SAR"],
        ["E014", "Lina George", "Operations", "Operations Director", "G12", 91000, "Exceeds", "Eligible", "Yes", "Low", "Saudi Arabia", "SAR"],
        ["E015", "Yasir Malik", "Sales", "Sales Specialist", "G8", 31200, "Does Not Meet", "Eligible", "No", "Low", "Saudi Arabia", "SAR"],
        ["E016", "Aisha Tariq", "IT", "AI Product Analyst", "G9", 26800, "Outstanding", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E017", "Hassan Omar", "Finance", "Senior Accountant", "G8", 26200, "Meets", "Eligible", "No", "Low", "Saudi Arabia", "SAR"],
        ["E018", "Noor Abdullah", "HR", "L&D Specialist", "G7", 15800, "Exceeds", "Eligible", "No", "Medium", "Saudi Arabia", "SAR"],
        ["E019", "Peter James", "Operations", "Project Manager", "G10", 47200, "Meets", "Eligible", "No", "Low", "Saudi Arabia", "SAR"],
        ["E020", "Maryam Saleh", "Sales", "Key Account Manager", "G10", 34500, "Outstanding", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E021", "Samir Khan", "IT", "IT Manager", "G10", 36000, "Exceeds", "Eligible", "Yes", "High", "Saudi Arabia", "SAR"],
        ["E022", "Rania Nabil", "Finance", "Budget Analyst", "G7", 14500, "Meets", "Eligible", "No", "Medium", "Saudi Arabia", "SAR"],
        ["E023", "Fahad Jamil", "Operations", "Maintenance Lead", "G8", 22750, "Partially Meets", "Eligible", "No", "Low", "Saudi Arabia", "SAR"],
        ["E024", "Leena Joseph", "HR", "Recruitment Lead", "G9", 32200, "Exceeds", "Eligible", "No", "Medium", "Saudi Arabia", "SAR"],
        ["E025", "Majed Anwar", "Sales", "Sales Director", "G12", 65500, "Meets", "Eligible", "Yes", "Medium", "Saudi Arabia", "SAR"],
    ]
    return pd.DataFrame(rows, columns=[
        "Employee_ID", "Employee_Name", "Department", "Job_Title", "Grade", "Current_Base_Salary",
        "Performance_Rating", "Eligibility", "Critical_Role", "Flight_Risk", "Country", "Currency"
    ])


def sample_salary_structure() -> pd.DataFrame:
    return pd.DataFrame([
        ["G6", 12000, 16000, 22000, "Professional"],
        ["G7", 15000, 20000, 27000, "Senior Professional"],
        ["G8", 18000, 24000, 33000, "Lead / Specialist"],
        ["G9", 22000, 30000, 40000, "Supervisor / Expert"],
        ["G10", 28000, 38000, 50000, "Manager"],
        ["G11", 35000, 48000, 65000, "Senior Manager"],
        ["G12", 45000, 62000, 85000, "Director"],
    ], columns=["Grade", "Minimum", "Midpoint", "Maximum", "Career_Level"])


def sample_merit_matrix() -> pd.DataFrame:
    rows = []
    values = {
        "Outstanding": {"Low": 0.07, "Mid": 0.06, "High": 0.04},
        "Exceeds": {"Low": 0.06, "Mid": 0.05, "High": 0.03},
        "Meets": {"Low": 0.04, "Mid": 0.03, "High": 0.02},
        "Partially Meets": {"Low": 0.01, "Mid": 0.00, "High": 0.00},
        "Does Not Meet": {"Low": 0.00, "Mid": 0.00, "High": 0.00},
    }
    for perf, zones in values.items():
        for zone, pct in zones.items():
            rows.append([perf, zone, pct])
    return pd.DataFrame(rows, columns=["Performance_Rating", "Compa_Zone", "Merit_Percent"])


def sample_market_benchmark() -> pd.DataFrame:
    titles = sample_employee_data()[["Job_Title", "Grade"]].drop_duplicates()
    structure = sample_salary_structure()[["Grade", "Midpoint"]]
    m = titles.merge(structure, on="Grade", how="left")
    m["Market_P25"] = (m["Midpoint"] * 0.88).round(0)
    m["Market_P50"] = (m["Midpoint"] * np.where(m["Job_Title"].str.contains("Cyber|Cloud|AI|Data", case=False), 1.15, 1.00)).round(0)
    m["Market_P75"] = (m["Market_P50"] * 1.15).round(0)
    m["Market_P90"] = (m["Market_P50"] * 1.30).round(0)
    m["Market_Source"] = "Demo Market Data"
    return m[["Job_Title", "Grade", "Market_P25", "Market_P50", "Market_P75", "Market_P90", "Market_Source"]]


def sample_promotions() -> pd.DataFrame:
    return pd.DataFrame([
        ["E004", "G9", "G10", 0.08, "Critical skill retention"],
        ["E006", "G8", "G9", 0.08, "Role expansion"],
        ["E018", "G7", "G8", 0.07, "Career progression"],
        ["E020", "G10", "G11", 0.10, "Expanded account leadership"],
        ["E024", "G9", "G10", 0.08, "Recruitment leadership scope"],
    ], columns=["Employee_ID", "Current_Grade", "Proposed_Grade", "Promotion_Percent", "Promotion_Reason"])


def sample_budget() -> pd.DataFrame:
    payroll = sample_employee_data().loc[sample_employee_data()["Eligibility"].eq("Eligible"), "Current_Base_Salary"].sum()
    return pd.DataFrame([{
        "Eligible_Payroll": payroll,
        "Merit_Budget_Percent": 0.04,
        "Promotion_Budget_Percent": 0.015,
        "Market_Correction_Budget_Percent": 0.010,
        "Currency": "SAR",
    }])


def create_workbook_bytes(blank: bool = False) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        if blank:
            sample_employee_data().head(0).to_excel(writer, sheet_name="Employee_Data", index=False)
            sample_salary_structure().head(0).to_excel(writer, sheet_name="Salary_Structure", index=False)
            sample_merit_matrix().head(0).to_excel(writer, sheet_name="Merit_Matrix", index=False)
            sample_budget().head(0).to_excel(writer, sheet_name="Budget_Assumptions", index=False)
            sample_market_benchmark().head(0).to_excel(writer, sheet_name="Market_Benchmark", index=False)
            sample_promotions().head(0).to_excel(writer, sheet_name="Promotion_Nominations", index=False)
        else:
            sample_employee_data().to_excel(writer, sheet_name="Employee_Data", index=False)
            sample_salary_structure().to_excel(writer, sheet_name="Salary_Structure", index=False)
            sample_merit_matrix().to_excel(writer, sheet_name="Merit_Matrix", index=False)
            sample_budget().to_excel(writer, sheet_name="Budget_Assumptions", index=False)
            sample_market_benchmark().to_excel(writer, sheet_name="Market_Benchmark", index=False)
            sample_promotions().to_excel(writer, sheet_name="Promotion_Nominations", index=False)
    output.seek(0)
    return output.getvalue()

# -------------------------------
# Loading and Validation
# -------------------------------

def read_workbook(file) -> Dict[str, pd.DataFrame]:
    xls = pd.ExcelFile(file)
    sheets = {}
    for name in xls.sheet_names:
        sheets[name] = pd.read_excel(xls, sheet_name=name)
    return sheets


def load_sample() -> Dict[str, pd.DataFrame]:
    return {
        "Employee_Data": sample_employee_data(),
        "Salary_Structure": sample_salary_structure(),
        "Merit_Matrix": sample_merit_matrix(),
        "Budget_Assumptions": sample_budget(),
        "Market_Benchmark": sample_market_benchmark(),
        "Promotion_Nominations": sample_promotions(),
    }


def normalize_percent(series: pd.Series) -> pd.Series:
    s = pd.to_numeric(series, errors="coerce").fillna(0)
    return np.where(s > 1, s / 100, s)


def validate_data(sheets: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    issues = []
    required_sheets = ["Employee_Data", "Salary_Structure", "Merit_Matrix", "Budget_Assumptions"]
    for sh in required_sheets:
        if sh not in sheets:
            issues.append([f"Missing sheet: {sh}", 1, "High", f"Add sheet named {sh}."])
    if issues:
        return pd.DataFrame(issues, columns=["Issue Type", "Count", "Severity", "Recommended Action"])

    emp = sheets["Employee_Data"].copy()
    ss = sheets["Salary_Structure"].copy()
    mm = sheets["Merit_Matrix"].copy()

    checks = [
        ("Missing Employee ID", emp.get("Employee_ID", pd.Series(dtype=object)).isna().sum(), "High", "Complete Employee_ID for every row."),
        ("Duplicate Employee ID", emp.get("Employee_ID", pd.Series(dtype=object)).duplicated().sum(), "High", "Remove duplicate employee records."),
        ("Missing Grade", emp.get("Grade", pd.Series(dtype=object)).isna().sum(), "High", "Map every employee to a valid grade."),
        ("Missing Current Salary", emp.get("Current_Base_Salary", pd.Series(dtype=float)).isna().sum(), "High", "Complete salary for every employee."),
        ("Missing Performance Rating", emp.get("Performance_Rating", pd.Series(dtype=object)).isna().sum(), "Medium", "Complete performance rating or mark ineligible."),
    ]
    for issue, count, sev, action in checks:
        if int(count) > 0:
            issues.append([issue, int(count), sev, action])

    if "Grade" in emp.columns and "Grade" in ss.columns:
        missing_grade = ~emp["Grade"].isin(ss["Grade"])
        c = int(missing_grade.sum())
        if c > 0:
            issues.append(["Employee grade not found in salary structure", c, "High", "Add missing grades to Salary_Structure."])

    if {"Performance_Rating"}.issubset(emp.columns) and {"Performance_Rating"}.issubset(mm.columns):
        missing_perf = ~emp["Performance_Rating"].isin(mm["Performance_Rating"].dropna().unique())
        c = int(missing_perf.sum())
        if c > 0:
            issues.append(["Performance rating not found in merit matrix", c, "Medium", "Align Performance_Rating values with Merit_Matrix."])

    if not issues:
        issues.append(["No blocking validation issues", 0, "Good", "Proceed to pay analysis."])
    return pd.DataFrame(issues, columns=["Issue Type", "Count", "Severity", "Recommended Action"])

# -------------------------------
# Calculation Engine
# -------------------------------

def calculate_analysis(sheets: Dict[str, pd.DataFrame]) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, dict]:
    emp = sheets["Employee_Data"].copy()
    ss = sheets["Salary_Structure"].copy()
    mm = sheets["Merit_Matrix"].copy()
    budget = sheets.get("Budget_Assumptions", sample_budget()).copy()
    market = sheets.get("Market_Benchmark", sample_market_benchmark()).copy()
    promos = sheets.get("Promotion_Nominations", sample_promotions()).copy()

    # Clean numeric columns
    for col in ["Current_Base_Salary"]:
        emp[col] = pd.to_numeric(emp[col], errors="coerce")
    for col in ["Minimum", "Midpoint", "Maximum"]:
        ss[col] = pd.to_numeric(ss[col], errors="coerce")
    mm["Merit_Percent"] = normalize_percent(mm["Merit_Percent"])

    df = emp.merge(ss, on="Grade", how="left")
    df["Compa_Ratio"] = df["Current_Base_Salary"] / df["Midpoint"]
    df["Range_Penetration"] = (df["Current_Base_Salary"] - df["Minimum"]) / (df["Maximum"] - df["Minimum"])
    df["Compa_Zone"] = np.select(
        [df["Compa_Ratio"] < 0.90, df["Compa_Ratio"].between(0.90, 1.10, inclusive="both"), df["Compa_Ratio"] > 1.10],
        ["Low", "Mid", "High"],
        default="Unmapped",
    )
    df["Salary_Status"] = np.select(
        [df["Current_Base_Salary"] < df["Minimum"], df["Current_Base_Salary"] > df["Maximum"]],
        ["Below Minimum", "Above Maximum"],
        default="Within Range",
    )

    df = df.merge(mm, on=["Performance_Rating", "Compa_Zone"], how="left")
    df["Merit_Percent"] = df["Merit_Percent"].fillna(0)
    df["Eligible_Flag"] = df["Eligibility"].astype(str).str.lower().eq("eligible")
    df["Merit_Increase_Amount"] = np.where(df["Eligible_Flag"], df["Current_Base_Salary"] * df["Merit_Percent"], 0)
    df["Salary_After_Merit"] = df["Current_Base_Salary"] + df["Merit_Increase_Amount"]
    df["New_Salary_Above_Max"] = df["Salary_After_Merit"] > df["Maximum"]
    df["Base_Increase_Allowed"] = np.where(
        df["New_Salary_Above_Max"],
        np.maximum(df["Maximum"] - df["Current_Base_Salary"], 0),
        df["Merit_Increase_Amount"],
    )
    df["Lump_Sum_Amount"] = np.maximum(df["Merit_Increase_Amount"] - df["Base_Increase_Allowed"], 0)
    df["Final_Base_Salary"] = df["Current_Base_Salary"] + df["Base_Increase_Allowed"]

    if not market.empty:
        for col in ["Market_P25", "Market_P50", "Market_P75", "Market_P90"]:
            if col in market.columns:
                market[col] = pd.to_numeric(market[col], errors="coerce")
        df = df.merge(market, on=["Job_Title", "Grade"], how="left")
        df["Market_Ratio_P50"] = df["Current_Base_Salary"] / df["Market_P50"]
        df["Market_Ratio_P75"] = df["Current_Base_Salary"] / df["Market_P75"]
        df["Gap_to_P50"] = df["Market_P50"] - df["Current_Base_Salary"]
        df["Gap_to_P75"] = df["Market_P75"] - df["Current_Base_Salary"]
    else:
        for col in ["Market_P25", "Market_P50", "Market_P75", "Market_P90", "Market_Ratio_P50", "Market_Ratio_P75", "Gap_to_P50", "Gap_to_P75"]:
            df[col] = np.nan

    # Promotion analysis
    promo_out = pd.DataFrame()
    if not promos.empty and "Employee_ID" in promos.columns:
        p = promos.copy()
        p["Promotion_Percent"] = normalize_percent(p.get("Promotion_Percent", pd.Series([0.08]*len(p))))
        base_cols = ["Employee_ID", "Employee_Name", "Department", "Job_Title", "Current_Base_Salary", "Grade"]
        p = p.merge(emp[base_cols], on="Employee_ID", how="left")
        new_range = ss.rename(columns={"Grade": "Proposed_Grade", "Minimum": "New_Minimum", "Midpoint": "New_Midpoint", "Maximum": "New_Maximum"})
        p = p.merge(new_range[["Proposed_Grade", "New_Minimum", "New_Midpoint", "New_Maximum"]], on="Proposed_Grade", how="left")
        p["Policy_Promotion_Increase"] = p["Current_Base_Salary"] * p["Promotion_Percent"]
        p["Minimum_Adjustment"] = np.maximum(p["New_Minimum"] - p["Current_Base_Salary"], 0)
        p["Promotion_Increase_Amount"] = np.maximum(p["Policy_Promotion_Increase"], p["Minimum_Adjustment"])
        p["Salary_After_Promotion"] = p["Current_Base_Salary"] + p["Promotion_Increase_Amount"]
        p["Promotion_Status"] = np.select(
            [p["Salary_After_Promotion"] < p["New_Minimum"], p["Salary_After_Promotion"] > p["New_Maximum"]],
            ["Below New Minimum", "Above New Maximum"],
            default="Within New Range",
        )
        promo_out = p

    # Risk flags
    risks = []
    for _, r in df.iterrows():
        def add_risk(risk_type, severity, reason, action):
            risks.append({
                "Employee_ID": r.get("Employee_ID"), "Employee_Name": r.get("Employee_Name"), "Department": r.get("Department"),
                "Grade": r.get("Grade"), "Risk_Type": risk_type, "Severity": severity, "Reason": reason, "Recommended_Action": action,
            })
        if r.get("Salary_Status") == "Below Minimum":
            add_risk("Below Minimum", "High", "Current salary is below the approved grade minimum.", "Review structural adjustment to bring salary to minimum.")
        if r.get("Salary_Status") == "Above Maximum":
            add_risk("Above Maximum", "Medium", "Current salary is above the approved grade maximum.", "Avoid further base increase; consider lump-sum/non-base reward.")
        if r.get("Performance_Rating") in ["Outstanding", "Exceeds"] and r.get("Compa_Ratio", 1) < 0.90:
            add_risk("High Performer Underpaid", "High", "High performance rating with compa-ratio below 90%.", "Prioritize merit or targeted retention adjustment.")
        if r.get("Performance_Rating") in ["Partially Meets", "Does Not Meet"] and r.get("Compa_Ratio", 0) > 1.05:
            add_risk("Low Performer Overpaid", "Medium", "Low performance rating with pay position above 105% of midpoint.", "Limit base salary movement and review performance plan.")
        if str(r.get("Critical_Role", "")).lower() == "yes" and r.get("Compa_Ratio", 1) < 0.90:
            add_risk("Critical Role Pay Risk", "High", "Critical role with compa-ratio below 90%.", "Review market competitiveness and retention risk.")
        if pd.notna(r.get("Market_Ratio_P50")) and r.get("Market_Ratio_P50") < 0.80:
            add_risk("Market Gap Above 20%", "High", "Current salary is more than 20% below market P50.", "Review targeted market correction subject to budget.")
        if bool(r.get("New_Salary_Above_Max", False)):
            add_risk("New Salary Above Maximum", "High", "Proposed merit movement pushes salary above range maximum.", "Cap base increase at maximum and convert excess to lump-sum.")
    risk_df = pd.DataFrame(risks)
    if risk_df.empty:
        risk_df = pd.DataFrame(columns=["Employee_ID", "Employee_Name", "Department", "Grade", "Risk_Type", "Severity", "Reason", "Recommended_Action"])

    # Summaries
    dept_summary = df.groupby("Department", dropna=False).agg(
        Headcount=("Employee_ID", "count"),
        Eligible_Headcount=("Eligible_Flag", "sum"),
        Avg_Salary=("Current_Base_Salary", "mean"),
        Avg_Compa_Ratio=("Compa_Ratio", "mean"),
        Avg_Market_Ratio_P50=("Market_Ratio_P50", "mean"),
        Avg_Merit_Percent=("Merit_Percent", "mean"),
        Proposed_Merit_Cost=("Merit_Increase_Amount", "sum"),
        Lump_Sum_Cost=("Lump_Sum_Amount", "sum"),
    ).reset_index()
    grade_summary = df.groupby("Grade", dropna=False).agg(
        Headcount=("Employee_ID", "count"),
        Avg_Salary=("Current_Base_Salary", "mean"),
        Avg_Compa_Ratio=("Compa_Ratio", "mean"),
        Below_Minimum=("Salary_Status", lambda s: (s == "Below Minimum").sum()),
        Above_Maximum=("Salary_Status", lambda s: (s == "Above Maximum").sum()),
        Proposed_Merit_Cost=("Merit_Increase_Amount", "sum"),
    ).reset_index()

    b = budget.iloc[0].to_dict() if not budget.empty else sample_budget().iloc[0].to_dict()
    eligible_payroll = float(b.get("Eligible_Payroll", 0) or 0)
    if eligible_payroll <= 0:
        eligible_payroll = df.loc[df["Eligible_Flag"], "Current_Base_Salary"].sum()
    merit_budget_pct = float(b.get("Merit_Budget_Percent", 0.04) or 0.04)
    if merit_budget_pct > 1:
        merit_budget_pct = merit_budget_pct / 100
    promotion_budget_pct = float(b.get("Promotion_Budget_Percent", 0.015) or 0.015)
    if promotion_budget_pct > 1:
        promotion_budget_pct = promotion_budget_pct / 100
    market_budget_pct = float(b.get("Market_Correction_Budget_Percent", 0.01) or 0.01)
    if market_budget_pct > 1:
        market_budget_pct = market_budget_pct / 100
    currency = b.get("Currency", "SAR")

    merit_budget = eligible_payroll * merit_budget_pct
    promotion_budget = eligible_payroll * promotion_budget_pct
    market_budget = eligible_payroll * market_budget_pct
    merit_cost = df["Merit_Increase_Amount"].sum()
    promo_cost = promo_out["Promotion_Increase_Amount"].sum() if not promo_out.empty else 0
    lump_sum_cost = df["Lump_Sum_Amount"].sum()
    total_cost = merit_cost + promo_cost + lump_sum_cost
    total_budget = merit_budget + promotion_budget + market_budget

    summary = {
        "currency": currency,
        "total_employees": int(df["Employee_ID"].nunique()),
        "eligible_employees": int(df["Eligible_Flag"].sum()),
        "eligible_payroll": eligible_payroll,
        "merit_budget": merit_budget,
        "promotion_budget": promotion_budget,
        "market_budget": market_budget,
        "total_budget": total_budget,
        "merit_cost": merit_cost,
        "promotion_cost": promo_cost,
        "lump_sum_cost": lump_sum_cost,
        "total_cost": total_cost,
        "budget_utilization": total_cost / total_budget if total_budget else 0,
        "below_minimum_count": int((df["Salary_Status"] == "Below Minimum").sum()),
        "above_maximum_count": int((df["Salary_Status"] == "Above Maximum").sum()),
        "high_performer_underpaid_count": int(((df["Performance_Rating"].isin(["Outstanding", "Exceeds"])) & (df["Compa_Ratio"] < 0.90)).sum()),
        "critical_role_risk_count": int(((df["Critical_Role"].astype(str).str.lower() == "yes") & (df["Compa_Ratio"] < 0.90)).sum()),
        "risk_count": int(len(risk_df)),
    }

    return df, dept_summary, grade_summary, risk_df, promo_out, summary


def scenario_table(df: pd.DataFrame, promo_out: pd.DataFrame, summary: dict) -> pd.DataFrame:
    base_merit = df["Merit_Increase_Amount"].sum()
    promo = promo_out["Promotion_Increase_Amount"].sum() if not promo_out.empty else 0
    critical_mask = df["Critical_Role"].astype(str).str.lower().eq("yes")
    high_perf_mask = df["Performance_Rating"].isin(["Outstanding", "Exceeds"])
    scenarios = []
    configs = [
        ("Base Case", 1.00, 1.00, "Medium", "Balanced plan based on current merit matrix."),
        ("Conservative", 0.85, 0.85, "High", "Lower cost but higher retention risk."),
        ("High Differentiation", 1.10, 1.00, "Low", "Higher differentiation for performance; may stretch budget."),
        ("Critical Talent Focus", 1.00, 1.00, "Low", "Prioritizes critical roles and high performers below midpoint."),
    ]
    for name, merit_mult, promo_mult, risk, note in configs:
        if name == "Critical Talent Focus":
            extra = (df.loc[critical_mask & high_perf_mask & (df["Compa_Ratio"] < 0.95), "Current_Base_Salary"] * 0.015).sum()
            merit_cost = base_merit + extra
        else:
            merit_cost = base_merit * merit_mult
        promo_cost = promo * promo_mult
        lump = df["Lump_Sum_Amount"].sum()
        total = merit_cost + promo_cost + lump
        util = total / summary["total_budget"] if summary["total_budget"] else 0
        scenarios.append([name, merit_cost, promo_cost, lump, total, util, risk, note])
    return pd.DataFrame(scenarios, columns=["Scenario", "Merit Cost", "Promotion Cost", "Lump-Sum Cost", "Total Cost", "Budget Utilization", "Talent Risk", "Recommendation"])

# -------------------------------
# Reports and Exports
# -------------------------------

def money(x, currency="SAR"):
    try:
        return f"{currency} {x:,.0f}"
    except Exception:
        return f"{currency} 0"


def pct(x):
    try:
        return f"{x:.1%}"
    except Exception:
        return "0.0%"


def generate_report(report_type: str, df: pd.DataFrame, risk_df: pd.DataFrame, promo_out: pd.DataFrame, summary: dict, scenarios: pd.DataFrame) -> str:
    c = summary["currency"]
    recommended_scenario = scenarios.sort_values("Budget Utilization", ascending=False).iloc[0]["Scenario"] if not scenarios.empty else "Base Case"
    if "CFO" in report_type:
        return f"""# CFO Affordability Note

The proposed compensation review has a total estimated reward cost of **{money(summary['total_cost'], c)}** against a total planning budget of **{money(summary['total_budget'], c)}**, resulting in budget utilization of **{pct(summary['budget_utilization'])}**.

## Cost Summary
- Eligible payroll: **{money(summary['eligible_payroll'], c)}**
- Merit cost: **{money(summary['merit_cost'], c)}**
- Promotion cost: **{money(summary['promotion_cost'], c)}**
- Lump-sum cost: **{money(summary['lump_sum_cost'], c)}**
- Total proposed cost: **{money(summary['total_cost'], c)}**

## Finance View
The plan should be reviewed against affordability, retention risk, and exception controls before final approval. Budget pressure is acceptable if utilization remains within approved thresholds and exception cases are formally documented.
"""
    if "NRC" in report_type:
        return f"""# NRC Approval Paper

## Background
Management has completed the annual compensation review using salary structure positioning, performance ratings, merit matrix logic, market benchmark signals, promotion nominations, and governance risk flags.

## Management Proposal
Management proposes a total reward movement cost of **{money(summary['total_cost'], c)}**, representing **{pct(summary['budget_utilization'])}** utilization of the total planned compensation review budget.

## Compensation Principles Applied
- Performance-linked differentiation
- Salary range governance through compa-ratio and range penetration
- Budget discipline and affordability review
- Risk review for below-minimum, above-maximum, and critical role cases

## Key Governance Flags
- Below minimum cases: **{summary['below_minimum_count']}**
- Above maximum cases: **{summary['above_maximum_count']}**
- High performer underpaid cases: **{summary['high_performer_underpaid_count']}**
- Critical role pay risk cases: **{summary['critical_role_risk_count']}**

## Decision Required
Approval is requested for the proposed annual salary review plan, subject to final HR and Finance validation of exceptions.

## Recommended Approval Wording
Management recommends approval of the proposed annual compensation review with a total estimated cost of **{money(summary['total_cost'], c)}**, aligned with the approved compensation philosophy, affordability envelope, and reward governance principles.
"""
    if "Manager" in report_type:
        return """# Manager Communication Pack

## Core Message
This year’s salary review considered performance rating, salary range position, approved budget, and company compensation guidelines.

## How to Explain the Decision
- Performance influenced the merit percentage.
- Salary range position influenced whether the increase was higher, moderate, or limited.
- Employees already high in the range may receive a lower base salary movement.
- Employees at or above range maximum may be considered for lump-sum treatment instead of additional base salary movement.

## Important Guidance
Managers should not compare employee outcomes with other employees. The conversation should focus on individual contribution, role, performance, salary range position, and future development.
"""
    if "HRBP" in report_type:
        top_depts = risk_df["Department"].value_counts().head(3).to_dict() if not risk_df.empty else {}
        return f"""# HRBP Department Report

HRBPs should focus on departments with the highest concentration of compensation risks. Current high-risk themes include below-minimum cases, high performers below midpoint, critical roles below competitive pay position, and employees whose proposed salary movement may exceed range maximum.

## Priority Departments
{top_depts if top_depts else 'No material department-level risk concentration identified.'}

## HRBP Actions
1. Review employee-level exceptions with department leaders.
2. Validate whether critical role flags align with current workforce priorities.
3. Prepare manager communication points before employee discussions.
4. Escalate any budget or governance exceptions to Total Rewards.
"""
    if "Consultant" in report_type:
        return f"""# Consultant Advisory Note

The compensation review shows a total proposed cost of **{money(summary['total_cost'], c)}** with budget utilization of **{pct(summary['budget_utilization'])}**. The analysis identifies **{summary['risk_count']}** total governance flags requiring review.

## Advisory View
The current plan is directionally sound if management intends to balance affordability with retention of high performers and critical roles. However, below-minimum cases and high performer underpaid cases should be reviewed before final approval.

## Recommended Consulting Actions
1. Validate salary structure alignment for grades with low average compa-ratio.
2. Prioritize targeted adjustments over across-the-board increases.
3. Use lump-sum treatment for employees at or above range maximum.
4. Present a scenario comparison to CHRO and CFO before final approval.
"""
    if "FAQ" in report_type:
        return """# Employee Salary Review FAQ

## What factors were considered?
Salary review decisions considered performance, salary range position, company budget, and compensation policy.

## Does everyone receive the same increase?
No. Increases may vary based on performance and current salary position within the approved range.

## Why might someone receive a lower increase?
Employees already positioned high in their salary range may receive a lower increase to maintain pay governance.

## What happens if someone is above the range maximum?
The company may consider a lump-sum or non-base reward instead of increasing base salary further.
"""
    # Default CHRO Summary
    return f"""# CHRO Salary Review Summary

The annual compensation review analyzed **{summary['total_employees']}** employees, including **{summary['eligible_employees']}** eligible employees. The proposed total reward movement is **{money(summary['total_cost'], c)}**, representing **{pct(summary['budget_utilization'])}** utilization of the planned reward budget.

## Key Findings
- Below minimum cases: **{summary['below_minimum_count']}**
- Above maximum cases: **{summary['above_maximum_count']}**
- High performer underpaid cases: **{summary['high_performer_underpaid_count']}**
- Critical role pay risk cases: **{summary['critical_role_risk_count']}**
- Total governance risk flags: **{summary['risk_count']}**

## Management Interpretation
The salary review is broadly structured around performance differentiation, pay range position, and affordability. Priority review should focus on high performers below midpoint, critical roles below competitive positioning, and employees at or above range maximum.

## Recommended Actions
1. Review all below-minimum cases for structural adjustment.
2. Prioritize high performers and critical roles below midpoint.
3. Apply lump-sum treatment where proposed increases exceed range maximum.
4. Validate total cost with Finance before final approval.
5. Prepare manager communication before employee discussions.

## Recommended Scenario
The scenario simulator should be reviewed with Finance and leadership. Current recommended planning view: **{recommended_scenario}**.
"""


def export_excel(df, dept, grade, risk_df, promo_out, scenarios, summary) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, sheet_name="Employee_Analysis", index=False)
        dept.to_excel(writer, sheet_name="Department_Summary", index=False)
        grade.to_excel(writer, sheet_name="Grade_Summary", index=False)
        risk_df.to_excel(writer, sheet_name="Risk_Flags", index=False)
        promo_out.to_excel(writer, sheet_name="Promotion_Planner", index=False)
        scenarios.to_excel(writer, sheet_name="Scenarios", index=False)
        pd.DataFrame([summary]).to_excel(writer, sheet_name="Budget_Summary", index=False)
    output.seek(0)
    return output.getvalue()


def export_docx(report_md: str) -> bytes:
    if Document is None:
        return b""
    doc = Document()
    doc.add_heading("RewardAI Professional Report", level=1)
    for line in report_md.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", ""), style="List Bullet")
        else:
            doc.add_paragraph(line.replace("**", ""))
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

# -------------------------------
# Streamlit App
# -------------------------------

if "sheets" not in st.session_state:
    st.session_state.sheets = load_sample()
if "project_saved" not in st.session_state:
    st.session_state.project_saved = False

with st.sidebar:
    st.header("Project Setup")
    company = st.text_input("Company Name", "ABC Manufacturing")
    country = st.text_input("Country", "Saudi Arabia")
    currency = st.text_input("Currency", "SAR")
    year = st.number_input("Salary Review Year", min_value=2024, max_value=2035, value=2026, step=1)
    pay_strategy = st.selectbox("Pay Strategy", ["Match Market", "Lead Market", "Target Critical Roles", "Conservative Budget Control"])
    report_type = st.selectbox("Report Type", [
        "CHRO Salary Review Summary", "CFO Affordability Note", "NRC Approval Paper", "HRBP Department Report",
        "Consultant Advisory Note", "Manager Communication Pack", "Employee FAQ"
    ])
    if st.button("Save Project Setup"):
        st.session_state.project_saved = True
        st.success("Project setup saved.")
    st.markdown("---")
    st.caption("RewardAI Professional | Local + GitHub/Streamlit-ready single-file build")

st.markdown(
    f"""
<div class="ra-hero">
<h1>RewardAI Professional — Compensation Review Platform</h1>
<p>Consultant-grade compensation intelligence for salary review, pay positioning, merit planning, market benchmarking, promotion analysis, budget scenarios, and executive reports.</p>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    """
<div class="privacy-box">
<b>Data privacy note:</b> RewardAI Professional is designed for compensation analysis and decision support. Salary and employee data should be handled according to company confidentiality, data privacy, and access control policies. For demos, use anonymized or sample data. This tool supports decisions and does not replace formal HR, Finance, Legal, or governance approval.
</div>
""",
    unsafe_allow_html=True,
)

st.subheader("1. Upload RewardAI Workbook")
st.write("Upload an Excel workbook with sheets: Employee_Data, Salary_Structure, Merit_Matrix, Budget_Assumptions, Market_Benchmark, Promotion_Nominations")

c1, c2, c3 = st.columns([2, 1, 1])
with c1:
    uploaded = st.file_uploader("Upload workbook", type=["xlsx"], label_visibility="collapsed")
    if uploaded is not None:
        try:
            st.session_state.sheets = read_workbook(uploaded)
            st.success("Workbook uploaded successfully.")
        except Exception as e:
            st.error(f"Could not read workbook: {e}")
with c2:
    st.download_button("Download Blank Template", data=create_workbook_bytes(blank=True), file_name="rewardai_professional_blank_template.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
with c3:
    st.download_button("Download Sample Workbook", data=create_workbook_bytes(blank=False), file_name="rewardai_professional_sample_workbook.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

if st.button("Load Sample Data", type="primary"):
    st.session_state.sheets = load_sample()
    st.success("Sample data loaded.")

validation = validate_data(st.session_state.sheets)
try:
    analysis, dept_summary, grade_summary, risk_df, promo_out, summary = calculate_analysis(st.session_state.sheets)
    scenarios = scenario_table(analysis, promo_out, summary)
except Exception as e:
    st.error(f"Analysis could not be calculated. Please check workbook structure. Error: {e}")
    st.stop()

main_tabs = st.tabs([
    "Executive Dashboard", "Data Validation", "Pay Positioning", "Market Benchmark", "Merit Planner",
    "Promotion Planner", "Lump-Sum Planner", "Scenario Simulator", "Budget & Governance", "AI Reports", "Export Center", "User Guide"
])

with main_tabs[0]:
    st.subheader("Executive Dashboard")
    k1, k2, k3, k4, k5 = st.columns(5)
    k1.metric("Total Employees", f"{summary['total_employees']:,}")
    k2.metric("Eligible Employees", f"{summary['eligible_employees']:,}")
    k3.metric("Current Payroll", money(summary["eligible_payroll"], summary["currency"]))
    k4.metric("Total Proposed Cost", money(summary["total_cost"], summary["currency"]))
    k5.metric("Budget Utilization", pct(summary["budget_utilization"]))

    r1, r2, r3, r4 = st.columns(4)
    r1.metric("Below Minimum", summary["below_minimum_count"])
    r2.metric("Above Maximum", summary["above_maximum_count"])
    r3.metric("High Performer Underpaid", summary["high_performer_underpaid_count"])
    r4.metric("Critical Role Risk", summary["critical_role_risk_count"])

    st.markdown("<div class='insight-box'><b>Executive Insight:</b> RewardAI links salary review decisions to salary range position, performance, market competitiveness, promotion impact, and affordability. Focus leadership review on high performers below midpoint, critical roles below competitive pay position, and employees at or above range maximum.</div>", unsafe_allow_html=True)

    a, b = st.columns(2)
    with a:
        fig = px.histogram(analysis, x="Compa_Ratio", nbins=12, title="Compa-Ratio Distribution")
        st.plotly_chart(fig, use_container_width=True)
    with b:
        fig = px.pie(analysis, names="Compa_Zone", title="Employees by Compa-Zone")
        st.plotly_chart(fig, use_container_width=True)

with main_tabs[1]:
    st.subheader("Data Validation")
    st.dataframe(validation, use_container_width=True, hide_index=True)
    if (validation["Severity"] == "High").any() and validation["Count"].sum() > 0:
        st.warning("High severity validation issues exist. Review before using outputs for decision-making.")
    else:
        st.success("No blocking validation issues found.")

with main_tabs[2]:
    st.subheader("Pay Positioning")
    col1, col2 = st.columns(2)
    with col1:
        fig = px.bar(grade_summary.sort_values("Grade"), x="Grade", y="Avg_Compa_Ratio", title="Average Compa-Ratio by Grade")
        st.plotly_chart(fig, use_container_width=True)
    with col2:
        fig = px.bar(dept_summary.sort_values("Avg_Compa_Ratio"), x="Department", y="Avg_Compa_Ratio", title="Average Compa-Ratio by Department")
        st.plotly_chart(fig, use_container_width=True)
    show_cols = ["Employee_ID", "Employee_Name", "Department", "Job_Title", "Grade", "Current_Base_Salary", "Minimum", "Midpoint", "Maximum", "Compa_Ratio", "Range_Penetration", "Compa_Zone", "Salary_Status"]
    st.dataframe(analysis[show_cols].sort_values("Compa_Ratio"), use_container_width=True, hide_index=True)

with main_tabs[3]:
    st.subheader("Market Benchmark Analyzer")
    market_cols = ["Employee_ID", "Employee_Name", "Department", "Job_Title", "Grade", "Current_Base_Salary", "Market_P50", "Market_P75", "Market_Ratio_P50", "Market_Ratio_P75", "Gap_to_P50", "Gap_to_P75"]
    st.dataframe(analysis[market_cols].sort_values("Market_Ratio_P50"), use_container_width=True, hide_index=True)
    col1, col2 = st.columns(2)
    with col1:
        fig = px.bar(dept_summary, x="Department", y="Avg_Market_Ratio_P50", title="Average Market Ratio P50 by Department")
        st.plotly_chart(fig, use_container_width=True)
    with col2:
        fig = px.scatter(analysis, x="Compa_Ratio", y="Market_Ratio_P50", color="Department", hover_name="Employee_Name", title="Internal Pay Position vs Market Ratio")
        st.plotly_chart(fig, use_container_width=True)

with main_tabs[4]:
    st.subheader("Merit Planner")
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Merit Cost", money(summary["merit_cost"], summary["currency"]))
    m2.metric("Average Merit %", pct(analysis["Merit_Percent"].mean()))
    m3.metric("Employees Receiving Increase", int((analysis["Merit_Increase_Amount"] > 0).sum()))
    m4.metric("Zero Increase", int((analysis["Merit_Increase_Amount"] == 0).sum()))
    merit_cols = ["Employee_ID", "Employee_Name", "Department", "Grade", "Current_Base_Salary", "Performance_Rating", "Compa_Zone", "Merit_Percent", "Merit_Increase_Amount", "Salary_After_Merit", "Final_Base_Salary", "Lump_Sum_Amount", "Salary_Status"]
    st.dataframe(analysis[merit_cols].sort_values("Merit_Increase_Amount", ascending=False), use_container_width=True, hide_index=True)

with main_tabs[5]:
    st.subheader("Promotion Planner")
    if promo_out.empty:
        st.info("No promotion nominations found.")
    else:
        p1, p2, p3 = st.columns(3)
        p1.metric("Promotion Nominations", len(promo_out))
        p2.metric("Promotion Cost", money(promo_out["Promotion_Increase_Amount"].sum(), summary["currency"]))
        p3.metric("Above New Max", int((promo_out["Promotion_Status"] == "Above New Maximum").sum()))
        st.dataframe(promo_out, use_container_width=True, hide_index=True)

with main_tabs[6]:
    st.subheader("Lump-Sum Planner")
    lump = analysis[(analysis["Lump_Sum_Amount"] > 0) | (analysis["Salary_Status"] == "Above Maximum")].copy()
    l1, l2, l3 = st.columns(3)
    l1.metric("Lump-Sum Candidates", len(lump))
    l2.metric("Lump-Sum Cost", money(summary["lump_sum_cost"], summary["currency"]))
    l3.metric("Above Maximum", summary["above_maximum_count"])
    if lump.empty:
        st.success("No lump-sum candidates identified under current merit logic.")
    else:
        cols = ["Employee_ID", "Employee_Name", "Department", "Grade", "Current_Base_Salary", "Maximum", "Merit_Increase_Amount", "Base_Increase_Allowed", "Lump_Sum_Amount", "Final_Base_Salary", "Salary_Status"]
        st.dataframe(lump[cols], use_container_width=True, hide_index=True)

with main_tabs[7]:
    st.subheader("Scenario Simulator")
    st.dataframe(scenarios, use_container_width=True, hide_index=True)
    fig = px.bar(scenarios, x="Scenario", y="Budget Utilization", color="Talent Risk", title="Scenario Budget Utilization")
    st.plotly_chart(fig, use_container_width=True)

with main_tabs[8]:
    st.subheader("Budget & Governance")
    b1, b2, b3, b4 = st.columns(4)
    b1.metric("Total Budget", money(summary["total_budget"], summary["currency"]))
    b2.metric("Total Cost", money(summary["total_cost"], summary["currency"]))
    b3.metric("Budget Variance", money(summary["total_budget"] - summary["total_cost"], summary["currency"]))
    b4.metric("Risk Flags", summary["risk_count"])
    st.progress(min(float(summary["budget_utilization"]), 1.0), text=f"Budget utilization: {pct(summary['budget_utilization'])}")
    st.write("### Risk Flags")
    st.dataframe(risk_df, use_container_width=True, hide_index=True)
    st.write("### Department Summary")
    st.dataframe(dept_summary, use_container_width=True, hide_index=True)
    st.write("### Grade Summary")
    st.dataframe(grade_summary, use_container_width=True, hide_index=True)

with main_tabs[9]:
    st.subheader("AI Reports")
    report_md = generate_report(report_type, analysis, risk_df, promo_out, summary, scenarios)
    st.markdown(report_md)
    st.download_button("Download Report as TXT", data=report_md.encode("utf-8"), file_name=f"rewardai_{report_type.lower().replace(' ', '_')}.txt", mime="text/plain", use_container_width=True)
    if Document is not None:
        docx_data = export_docx(report_md)
        st.download_button("Download Report as Word", data=docx_data, file_name=f"rewardai_{report_type.lower().replace(' ', '_')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    else:
        st.info("python-docx is not available, so Word export is disabled.")

with main_tabs[10]:
    st.subheader("Export Center")
    excel_data = export_excel(analysis, dept_summary, grade_summary, risk_df, promo_out, scenarios, summary)
    st.download_button("Download Full Excel Analysis", data=excel_data, file_name="rewardai_professional_full_analysis.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    st.download_button("Download Employee Recommendations CSV", data=analysis.to_csv(index=False).encode("utf-8"), file_name="employee_recommendations.csv", mime="text/csv", use_container_width=True)
    st.download_button("Download Risk Flags CSV", data=risk_df.to_csv(index=False).encode("utf-8"), file_name="risk_flags.csv", mime="text/csv", use_container_width=True)
    st.download_button("Download Budget Summary CSV", data=pd.DataFrame([summary]).to_csv(index=False).encode("utf-8"), file_name="budget_summary.csv", mime="text/csv", use_container_width=True)

with main_tabs[11]:
    st.subheader("User Guide")
    st.markdown(
        """
### How to use RewardAI Professional

1. **Load sample data** or upload your own Excel workbook.
2. Your workbook should include these sheets: `Employee_Data`, `Salary_Structure`, `Merit_Matrix`, `Budget_Assumptions`, `Market_Benchmark`, and `Promotion_Nominations`.
3. Go to **Data Validation** and review any issues.
4. Use **Pay Positioning** to review compa-ratio and range penetration.
5. Use **Market Benchmark** to compare internal pay with P50/P75.
6. Use **Merit Planner** to review recommended merit increases.
7. Use **Promotion Planner** and **Lump-Sum Planner** for professional compensation decisions.
8. Use **Scenario Simulator** to compare cost and risk options.
9. Generate CHRO, CFO, NRC, HRBP, Consultant, Manager, or Employee FAQ reports.
10. Export the full Excel analysis.

### Key calculations
- Compa-Ratio = Current Base Salary / Midpoint
- Range Penetration = (Current Salary - Minimum) / (Maximum - Minimum)
- Merit Amount = Current Salary × Merit Percent
- Salary After Merit = Current Salary + Merit Amount
- Market Ratio P50 = Current Salary / Market P50
- Promotion Increase = max(Current Salary × Promotion %, New Grade Minimum - Current Salary)

### GitHub deployment note
This version is a **single-file GitHub-ready build**. Upload only:
- `app.py`
- `requirements.txt`
- `README.md`

Then deploy on Streamlit Cloud using `app.py` as the main file path.
"""
    )
